import os
import io
from typing import Optional
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

# Import CrewAI LLM and our custom crew runner
from crewai import LLM
from crew import ResumeBuilderCrew

# Document generation helpers
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import docx
from docx.shared import Pt, RGBColor

load_dotenv()

app = FastAPI(title="AI Resume Builder API", version="1.0.0")

# Enable CORS for React frontend integration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Helper: Parse inline markdown tags (bold/italics) for PDF generation
def parse_markdown_inlines(text: str) -> str:
    parts = text.split('**')
    new_parts = []
    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            new_parts.append(f"<b>{part}</b>")
        else:
            sub_parts = part.split('*')
            sub_new = []
            for s_idx, s_part in enumerate(sub_parts):
                if s_idx % 2 == 1:
                    sub_new.append(f"<i>{s_part}</i>")
                else:
                    sub_new.append(s_part)
            new_parts.append("".join(sub_new))
    return "".join(new_parts)

# Helper: Export Markdown resume text to PDF using ReportLab
def markdown_to_pdf(md_content: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=54, 
        leftMargin=54,
        topMargin=54, 
        bottomMargin=54
    )
    styles = getSampleStyleSheet()
    
    # Define custom styling palette
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=12,
        alignment=1 # Center
    )
    
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=10,
        spaceAfter=5,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#334155'),
        spaceBefore=7,
        spaceAfter=3,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#334155'),
        spaceAfter=3
    )

    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#334155'),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=2.5
    )

    story = []
    lines = md_content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 4))
            continue
        
        # Simple Markdown parsing mapping
        if stripped.startswith('# '):
            text = stripped[2:].replace('**', '').replace('*', '')
            if not story:
                story.append(Paragraph(text, title_style))
            else:
                story.append(Paragraph(text, h1_style))
        elif stripped.startswith('## '):
            text = stripped[3:].replace('**', '').replace('*', '')
            story.append(Paragraph(text, h1_style))
        elif stripped.startswith('### '):
            text = stripped[4:].replace('**', '').replace('*', '')
            story.append(Paragraph(text, h2_style))
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            text = parse_markdown_inlines(text)
            story.append(Paragraph(f"&bull; {text}", bullet_style))
        else:
            text = parse_markdown_inlines(stripped)
            story.append(Paragraph(text, body_style))
            
    doc.build(story)
    buffer.seek(0)
    return buffer

# Helper: Parse inline markdown tags (bold/italics) for Word generation
def parse_markdown_docx_inlines(paragraph, text: str):
    parts = text.split('**')
    for idx, part in enumerate(parts):
        is_bold = (idx % 2 == 1)
        sub_parts = part.split('*')
        for s_idx, s_part in enumerate(sub_parts):
            is_italic = (s_idx % 2 == 1)
            run = paragraph.add_run(s_part)
            run.bold = is_bold
            run.italic = is_italic

# Helper: Export Markdown resume text to DOCX
def markdown_to_docx(md_content: str) -> io.BytesIO:
    doc = docx.Document()
    
    # Configure default Arial 10.5 font style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    lines = md_content.split('\n')
    is_first_header = True
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        if stripped.startswith('# '):
            text = stripped[2:].replace('**', '').replace('*', '')
            if is_first_header:
                p = doc.add_paragraph()
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                is_first_header = False
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif stripped.startswith('## '):
            text = stripped[3:].replace('**', '').replace('*', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif stripped.startswith('### '):
            text = stripped[4:].replace('**', '').replace('*', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_markdown_docx_inlines(p, text)
        else:
            p = doc.add_paragraph()
            parse_markdown_docx_inlines(p, stripped)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# LLM Configuration Resolver
def resolve_llm(provider: str, api_key: Optional[str]) -> LLM:
    if provider == "gemini":
        key = api_key or os.environ.get("GEMINI_API_KEY")
        if not key:
            raise HTTPException(
                status_code=400, 
                detail="Gemini API Key is missing. Please check your config or environment variables."
            )
        # Using LiteLLM provider tag for Gemini 3.1 Flash Lite
        return LLM(model="gemini/gemini-3.1-flash-lite", api_key=key)
    elif provider == "openai":
        key = api_key or os.environ.get("OPENAI_API_KEY")
        if not key:
            raise HTTPException(
                status_code=400, 
                detail="OpenAI API Key is missing. Please check your config or environment variables."
            )
        return LLM(model="openai/gpt-4o-mini", api_key=key)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported LLM provider: {provider}")

# Mock Job Search Database
MOCK_JOBS = [
    {
        "id": 1,
        "title": "Junior Full Stack Developer",
        "company": "InnovateTech India",
        "location": "Bengaluru, Karnataka",
        "type": "Hybrid",
        "experience": "Junior",
        "description": "Work with React.js, FastAPI backend, and PostgreSQL to deliver stellar SaaS features. Docker and AWS experience is a plus.",
        "skills": ["react", "python", "fastapi", "postgresql", "docker", "aws", "javascript"],
        "link": "https://www.linkedin.com/jobs/view/mock-1"
    },
    {
        "id": 2,
        "title": "React.js Developer",
        "company": "Razorpay",
        "location": "Remote (India)",
        "type": "Remote",
        "experience": "Entry Level",
        "description": "Design responsive, high-performance UI components using React, TailwindCSS, and JavaScript (ES6+). Must be highly skilled in visual design.",
        "skills": ["react", "javascript", "css", "html", "tailwind"],
        "link": "https://www.linkedin.com/jobs/view/mock-2"
    },
    {
        "id": 3,
        "title": "Software Engineer Intern",
        "company": "Flipkart",
        "location": "Bengaluru, Karnataka",
        "type": "On-site",
        "experience": "Internship",
        "description": "Assist the backend team in designing scalable endpoints using FastAPI and deploying containerized microservices on AWS EC2.",
        "skills": ["python", "fastapi", "docker", "aws", "git"],
        "link": "https://www.linkedin.com/jobs/view/mock-3"
    },
    {
        "id": 4,
        "title": "Python Backend Engineer",
        "company": "Infosys Technologies",
        "location": "Hyderabad, Telangana",
        "type": "Hybrid",
        "experience": "Mid Level",
        "description": "Build high-speed data integration servers, REST APIs, and pipelines using Python, FastAPI, PostgreSQL, and AWS S3.",
        "skills": ["python", "fastapi", "postgresql", "aws", "sql"],
        "link": "https://www.linkedin.com/jobs/view/mock-4"
    },
    {
        "id": 5,
        "title": "Senior React Architect",
        "company": "TCS Digital",
        "location": "Remote (India)",
        "type": "Remote",
        "experience": "Senior Level",
        "description": "Lead design and orchestration of micro-frontend applications with complex state management using React.js and TypeScript.",
        "skills": ["react", "typescript", "javascript", "css"],
        "link": "https://www.linkedin.com/jobs/view/mock-5"
    },
    {
        "id": 6,
        "title": "Junior Python Developer",
        "company": "Wipro Analytics",
        "location": "Mumbai, Maharashtra",
        "type": "On-site",
        "experience": "Junior",
        "description": "Integrate algorithmic trading features and build lightweight API dashboards using Python, SQLite, and Flask.",
        "skills": ["python", "flask", "sqlite", "git"],
        "link": "https://www.linkedin.com/jobs/view/mock-6"
    }
]

# Pydantic schemas for Export Route
class ExportRequest(BaseModel):
    content: str
    format: str  # "pdf" or "docx"
    filename: Optional[str] = "export"

@app.get("/")
def read_root():
    return {"status": "success", "message": "AI Resume Builder API is online"}

@app.post("/api/generate-resume")
def generate_resume(
    student_profile: UploadFile = File(None),
    student_profile_text: Optional[str] = Form(None),
    job_description: UploadFile = File(None),
    job_description_text: Optional[str] = Form(None),
    provider: str = Form("gemini"),
    api_key: Optional[str] = Form(None)
):
    # Parse inputs (prefer files if uploaded, fallback to text fields)
    profile_content = ""
    if student_profile:
        profile_content = student_profile.file.read().decode("utf-8")
    elif student_profile_text:
        profile_content = student_profile_text
    else:
        raise HTTPException(status_code=400, detail="Student Profile content is required.")

    jd_content = ""
    if job_description:
        jd_content = job_description.file.read().decode("utf-8")
    elif job_description_text:
        jd_content = job_description_text
    else:
        raise HTTPException(status_code=400, detail="Job Description content is required.")

    # Initialize the LLM instance
    llm = resolve_llm(provider, api_key)

    try:
        # Run CrewAI workflow
        crew_runner = ResumeBuilderCrew(llm)
        crew_result = crew_runner.run(profile_content, jd_content, run_bonus=False)

        # Save generated outputs to /output directory (as required by Phase 1)
        os.makedirs("output", exist_ok=True)
        with open("output/resume.md", "w", encoding="utf-8") as f:
            f.write(crew_result["resume"])
        with open("output/ats_report.md", "w", encoding="utf-8") as f:
            f.write(crew_result["ats_report"])
        with open("output/improvement_plan.md", "w", encoding="utf-8") as f:
            f.write(crew_result["improvement_plan"])

        return {
            "status": "success",
            "data": crew_result
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

@app.post("/api/generate-bonus")
def generate_bonus(
    student_profile: Optional[str] = Form(None),
    job_description: Optional[str] = Form(None),
    provider: str = Form("gemini"),
    api_key: Optional[str] = Form(None)
):
    if not student_profile or not job_description:
        raise HTTPException(status_code=400, detail="Both Student Profile and Job Description are required.")

    llm = resolve_llm(provider, api_key)

    try:
        crew_runner = ResumeBuilderCrew(llm)
        crew_result = crew_runner.run(student_profile, job_description, run_bonus=True)

        return {
            "status": "success",
            "data": {
                "cover_letter": crew_result.get("cover_letter", ""),
                "linkedin_about": crew_result.get("linkedin_about", ""),
                "interview_questions": crew_result.get("interview_questions", "")
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Bonus generation failed: {str(e)}")

@app.get("/api/jobs")
def get_jobs(
    query: Optional[str] = None,
    location: Optional[str] = None,
    experience: Optional[str] = None
):
    results = []
    search_keywords = query.lower().split() if query else []
    
    for job in MOCK_JOBS:
        # Filter by Location
        if location and location.lower() != "all":
            if location.lower() not in job["location"].lower() and location.lower() not in job["type"].lower():
                continue
                
        # Filter by Experience
        if experience and experience.lower() != "all":
            # Match levels
            if experience.lower() not in job["experience"].lower():
                continue
        
        # Calculate matching score based on keyword overlap in title, skills, description
        score = 50  # Base score
        if query:
            match_count = 0
            title_lower = job["title"].lower()
            desc_lower = job["description"].lower()
            
            for word in search_keywords:
                if word in title_lower:
                    match_count += 3
                if word in desc_lower:
                    match_count += 1
                for skill in job["skills"]:
                    if word in skill:
                        match_count += 2
                        
            score += min(match_count * 8, 48)  # Cap score contribution
        else:
            # Random slight variance for presentation if search is empty
            score += (job["id"] * 7) % 25
            
        job_result = job.copy()
        job_result["match_score"] = score
        results.append(job_result)

    # Sort results by relevance/match score descending
    results.sort(key=lambda x: x["match_score"], reverse=True)
    return results

@app.post("/api/export")
def export_file(request: ExportRequest):
    content = request.content
    fmt = request.format.lower()
    filename = request.filename or "resume"
    
    if fmt == "pdf":
        pdf_buffer = markdown_to_pdf(content)
        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
        )
    elif fmt == "docx":
        docx_buffer = markdown_to_docx(content)
        return Response(
            content=docx_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid export format. Supported formats are 'pdf' and 'docx'.")

if __name__ == "__main__":
    import uvicorn
    # Start on port 8000
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
